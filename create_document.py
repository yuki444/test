"""Word文書作成ユーティリティ"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import sys


def create_document(content_data: dict, output_path: str = "output.docx") -> str:
    """
    コンテンツデータからWord文書を生成する。

    content_data: {
        "title": "文書タイトル",
        "sections": [
            {
                "heading": "見出し",
                "level": 1,
                "paragraphs": ["段落1", "段落2"]
            },
            ...
        ]
    }
    """
    doc = Document()

    if "title" in content_data:
        doc.add_heading(content_data["title"], level=0)

    for section in content_data.get("sections", []):
        level = section.get("level", 1)
        if "heading" in section:
            doc.add_heading(section["heading"], level=level)

        for para_text in section.get("paragraphs", []):
            if para_text.startswith("- "):
                doc.add_paragraph(para_text[2:], style="List Bullet")
            elif para_text.startswith("1. ") or para_text[0:2] in ["1.", "2.", "3."]:
                doc.add_paragraph(para_text.split(". ", 1)[-1], style="List Number")
            else:
                doc.add_paragraph(para_text)

        if section.get("table"):
            table_data = section["table"]
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"
            for r, row in enumerate(table_data):
                for c, cell_text in enumerate(row):
                    table.cell(r, c).text = cell_text

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        sample_content = {
            "title": "サンプル文書",
            "sections": [
                {
                    "heading": "はじめに",
                    "level": 1,
                    "paragraphs": ["この文書はサンプルです。", "python-docxを使って生成されました。"]
                },
                {
                    "heading": "箇条書きの例",
                    "level": 1,
                    "paragraphs": ["- 項目1", "- 項目2", "- 項目3"]
                },
                {
                    "heading": "表の例",
                    "level": 1,
                    "paragraphs": [],
                    "table": [
                        ["名前", "役割", "備考"],
                        ["田中", "開発者", "バックエンド"],
                        ["鈴木", "デザイナー", "UI/UX"]
                    ]
                }
            ]
        }
        output = create_document(sample_content, "sample.docx")
        print(f"サンプル文書を作成しました: {output}")
    else:
        json_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else "output.docx"
        with open(json_path, "r", encoding="utf-8") as f:
            content_data = json.load(f)
        output = create_document(content_data, output_path)
        print(f"文書を作成しました: {output}")


if __name__ == "__main__":
    main()
